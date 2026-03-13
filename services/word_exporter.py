from __future__ import annotations

from io import BytesIO
from typing import Any, Literal

from docx import Document

ExportMode = Literal["draft", "oa", "compare", "polish"]


def _to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (dict, list)):
        import json

        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value)


def _pick_final_text(mode: str, data: dict[str, Any]) -> str:
    if mode == "draft":
        return _to_text(data.get("specification")) or _to_text(data.get("claims"))
    if mode == "oa":
        return _to_text(data.get("final_reply_text")) or _to_text(data.get("argument_draft"))
    if mode == "compare":
        return (
            _to_text(data.get("final_compare_report"))
            or _to_text(data.get("risk_assessment_report"))
            or _to_text(data.get("amendment_suggestions"))
        )
    return (
        _to_text(data.get("polish_final_package"))
        or _to_text(data.get("optimized_claims_text"))
        or _to_text(data.get("optimized_specification_text"))
    )


def build_export_docx(mode: str, session_id: str, data: dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading(f"{mode.upper()} 导出文档", level=1)
    doc.add_paragraph(f"Session ID: {session_id}")
    doc.add_paragraph("")

    final_text = _pick_final_text(mode, data)
    if final_text:
        doc.add_heading("最终输出", level=2)
        doc.add_paragraph(final_text)

    doc.add_heading("节点输出快照", level=2)
    for key, value in data.items():
        doc.add_heading(str(key), level=3)
        text = _to_text(value)
        doc.add_paragraph(text if text else "(空)")

    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()

